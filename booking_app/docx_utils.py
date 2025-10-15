"""Utility helpers for working with DOCX templates inside the admin UI."""

from __future__ import annotations

from io import BytesIO
from typing import Dict

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document


def _normalize_text(value: str) -> str:
    """Return a whitespace-normalized text segment suitable for DOCX runs."""

    return value.replace("\xa0", " ")


def html_to_docx(html: str) -> BytesIO:
    """Convert a snippet of editable HTML into a DOCX binary stream.

    The admin contract template editor works with a browser-based, contenteditable
    representation of the document. This helper translates the HTML that the
    browser emits into a reasonably faithful DOCX document using python-docx.

    Only a subset of HTML is supported—paragraphs, headings, lists, tables, and
    common inline formatting (bold, italic, underline). Unsupported nodes are
    flattened into paragraphs so the resulting template always contains the
    merged textual content and placeholders the user positioned.
    """

    soup = BeautifulSoup(html, "html.parser")
    document = Document()

    root = soup.body or soup

    has_content = False

    block_names = {"p", "div", "table", "ul", "ol", "h1", "h2", "h3", "h4", "h5", "h6"}

    def append_inline(paragraph, node: Tag | NavigableString, styles: Dict[str, bool] | None = None):
        nonlocal has_content

        styles = styles.copy() if styles else {}

        if isinstance(node, NavigableString):
            text = _normalize_text(str(node))
            if text:
                run = paragraph.add_run(text)
                if styles.get("bold"):
                    run.bold = True
                if styles.get("italic"):
                    run.italic = True
                if styles.get("underline"):
                    run.underline = True
                has_content = True
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        if name in {"strong", "b"}:
            styles["bold"] = True
        elif name in {"em", "i"}:
            styles["italic"] = True
        elif name in {"u", "ins"}:
            styles["underline"] = True

        if name == "br":
            paragraph.add_run().add_break()
            has_content = True
            return

        if name == "span":
            inline_styles = node.attrs.get("style", "")
            if "font-weight" in inline_styles and "bold" in inline_styles:
                styles["bold"] = True
            if "font-style" in inline_styles and "italic" in inline_styles:
                styles["italic"] = True
            if "text-decoration" in inline_styles and "underline" in inline_styles:
                styles["underline"] = True

        for child in node.children:
            append_inline(paragraph, child, styles)

    def add_paragraph_from(node: Tag, style: str | None = None):
        paragraph = document.add_paragraph()
        if style:
            paragraph.style = style
        append_inline(paragraph, node)

    def process_list(list_node: Tag, ordered: bool):
        list_style = "List Number" if ordered else "List Bullet"
        for item in list_node.find_all("li", recursive=False):
            paragraph = document.add_paragraph(style=list_style)
            append_inline(paragraph, item)
            # Support nested lists
            for nested in item.find_all(["ul", "ol"], recursive=False):
                process_list(nested, nested.name.lower() == "ol")

    def clear_cell(cell):
        for paragraph in list(cell.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)

    def process_table(table_node: Tag):
        rows = table_node.find_all("tr", recursive=False)
        if not rows:
            return

        max_cols = max(
            (len(row.find_all(["td", "th"], recursive=False)) for row in rows),
            default=0,
        )
        if max_cols == 0:
            return

        table = document.add_table(rows=len(rows), cols=max_cols)

        for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"], recursive=False)
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                clear_cell(cell)
                paragraph = cell.add_paragraph()

                if c_idx < len(cells):
                    cell_node = cells[c_idx]
                    styles = {"bold": cell_node.name.lower() == "th"}
                    append_inline(paragraph, cell_node, styles)

    def process_block(node):
        nonlocal has_content

        if isinstance(node, NavigableString):
            text = _normalize_text(str(node))
            if text.strip():
                document.add_paragraph(text)
                has_content = True
            return

        if not isinstance(node, Tag):
            return

        name = (node.name or "").lower()

        if name in {"script", "style"}:
            return

        if name == "p":
            add_paragraph_from(node)
            return

        if name == "div":
            if any(isinstance(child, Tag) and (child.name or "").lower() in block_names for child in node.children):
                for child in node.children:
                    process_block(child)
            else:
                add_paragraph_from(node)
            return

        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            add_paragraph_from(node, style=name.upper())
            return

        if name in {"ul", "ol"}:
            process_list(node, ordered=name == "ol")
            return

        if name == "table":
            process_table(node)
            return

        if name == "br":
            document.add_paragraph()
            return

        for child in node.children:
            process_block(child)

    for child in root.children:
        process_block(child)

    if not has_content:
        raise ValueError("The contract template cannot be empty.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer